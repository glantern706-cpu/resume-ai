import os
import logging
import docx
import tempfile

class FileProcessor:
    def __init__(self):
        """Initialize file processor"""
        pass
    
    def extract_text(self, file_path):
        """
        Extract text from PDF or DOCX files
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise ValueError("PDF is password protected and cannot be processed")
                
                # Extract text from all pages
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if not text.strip():
                    raise ValueError("No readable text found in PDF")
                
                return text.strip()
                
        except Exception as e:
            logging.error(f"Error reading PDF: {str(e)}")
            raise ValueError(f"Could not read PDF file: {str(e)}")
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No readable text found in document")
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"Error reading DOCX: {str(e)}")
            raise ValueError(f"Could not read DOCX file: {str(e)}")
    
    def validate_file(self, file_path, max_size_mb=16):
        """
        Validate uploaded file
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise ValueError("File does not exist")
            
            # Check file size
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                raise ValueError(f"File too large. Maximum size is {max_size_mb}MB")
            
            # Check file extension
            file_extension = os.path.splitext(file_path)[1].lower()
            allowed_extensions = ['.pdf', '.docx', '.doc']
            
            if file_extension not in allowed_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            return True
            
        except Exception as e:
            logging.error(f"File validation error: {str(e)}")
            raise
